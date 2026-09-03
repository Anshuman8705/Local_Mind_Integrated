"""Black-box system test: every portal, endpoint and error path over real HTTP.

    # terminal 1
    python scripts/fake_ollama.py                      # or a real Ollama with qwen3:1.7b
    # terminal 2 (fresh database)
    DJANGO_DEBUG=true python manage.py migrate
    DJANGO_DEBUG=true python manage.py bootstrap_admin --email root@localmind.test
    DJANGO_DEBUG=true python manage.py runserver 127.0.0.1:8011
    # terminal 3
    python scripts/system_test.py [base_url] [--fake-ollama http://127.0.0.1:11434]

Each check prints PASS/FAIL with the step name; the exit code is the number
of failures. The script creates its own subject, users and content, so it can
run on a database that already holds data, and cleans nothing up (everything
it makes is prefixed SYS-).
"""
from __future__ import annotations

import io
import json
import mimetypes
import sys
import time
import urllib.error
import urllib.parse
import urllib.request
import uuid

BASE = "http://127.0.0.1:8011"
FAKE = None
for i, arg in enumerate(sys.argv[1:], 1):
    if arg == "--fake-ollama":
        FAKE = sys.argv[i + 1]
    elif not arg.startswith("http") or FAKE == arg:
        continue
    else:
        BASE = arg
API = BASE.rstrip("/") + "/api"

INITIAL = "Welcome@LocalMind1"
STRONG = "Str0ng!Passw0rd#2026"
RUN = uuid.uuid4().hex[:6]

passed = failed = 0


def check(name, cond, info=""):
    global passed, failed
    if cond:
        passed += 1
        print(f"  PASS  {name}")
    else:
        failed += 1
        print(f"  FAIL  {name}  {info}")
    return cond


def section(title):
    print(f"\n== {title}")


def call(method, path, body=None, tok=None, files=None, raw=False):
    headers = {}
    if tok:
        headers["Authorization"] = f"Bearer {tok}"
    if files:
        boundary = "----sys" + uuid.uuid4().hex
        buf = io.BytesIO()
        for key, value in (body or {}).items():
            buf.write(f"--{boundary}\r\nContent-Disposition: form-data; name=\"{key}\"\r\n\r\n{value}\r\n".encode())
        for key, (name, content) in files.items():
            ctype = mimetypes.guess_type(name)[0] or "application/octet-stream"
            buf.write(f"--{boundary}\r\nContent-Disposition: form-data; name=\"{key}\"; filename=\"{name}\"\r\nContent-Type: {ctype}\r\n\r\n".encode())
            buf.write(content)
            buf.write(b"\r\n")
        buf.write(f"--{boundary}--\r\n".encode())
        data = buf.getvalue()
        headers["Content-Type"] = f"multipart/form-data; boundary={boundary}"
    else:
        data = json.dumps(body).encode() if body is not None else None
        headers["Content-Type"] = "application/json"
    if "?" in path:
        base, qs = path.split("?", 1)
        path = base + "?" + urllib.parse.quote(qs, safe="=&")
    req = urllib.request.Request(API + path, data=data, headers=headers, method=method)
    try:
        r = urllib.request.urlopen(req, timeout=180)
        text = r.read()
        return r.status, (text if raw else (json.loads(text) if text else {}))
    except urllib.error.HTTPError as e:
        text = e.read()
        try:
            return e.code, json.loads(text) if text else {}
        except ValueError:
            return e.code, {"raw": text[:200].decode(errors="replace")}


def rows(d):
    return d["results"] if isinstance(d, dict) and "results" in d else d


def err(d):
    return (d.get("error") or {}).get("code") if isinstance(d, dict) else None


def fake_control(**kw):
    if not FAKE:
        return
    req = urllib.request.Request(FAKE.rstrip("/") + "/_control", data=json.dumps(kw).encode(),
                                 headers={"Content-Type": "application/json"}, method="POST")
    urllib.request.urlopen(req, timeout=5).read()


def make_docx():
    from docx import Document
    d = Document()
    d.add_heading("Operating Systems", level=1)
    d.add_paragraph("An operating system manages hardware and provides services to programs. " * 8)
    d.add_heading("Process Management", level=2)
    d.add_paragraph("Processes are programs in execution. Each process has its own address space and state. " * 10)
    d.add_heading("Scheduling", level=2)
    d.add_paragraph("Round robin gives each process a time slice. Priority scheduling runs the highest priority first. " * 10)
    d.add_heading("Memory Management", level=2)
    d.add_paragraph("Paging divides memory into fixed frames; segmentation into variable segments. " * 10)
    d.add_heading("Networks", level=1)
    d.add_heading("Transport Layer", level=2)
    d.add_paragraph("TCP is reliable and ordered; UDP is best effort and fast. " * 10)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def make_xlsx(headers, rows_):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.append(headers)
    for r in rows_:
        ws.append(r)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def login(role, email, password):
    return call("POST", f"/auth/login/{role}/", {"email": email, "password": password})


def login_and_settle(role, email):
    """Log in with whichever password the account currently has; change the initial one."""
    s, d = login(role, email, INITIAL)
    if s == 200 and d.get("must_change_password"):
        s2, d2 = call("POST", "/auth/password/change/", {"current_password": INITIAL, "new_password": STRONG}, tok=d["access"])
        check(f"{role} forced password change", s2 == 200 and d2.get("must_change_password") is False, d2)
        return d2["access"], d2["refresh"], d.get("session_id")
    if s != 200:
        s, d = login(role, email, STRONG)
    check(f"{role} login {email}", s == 200, d)
    return d["access"], d["refresh"], d.get("session_id")


fake_control(offline=0, fail_next=0)

# ============================================================ health ==========
section("Health")
s, d = call("GET", "/health/")
check("health responds", s == 200 and d.get("database") == "ok", d)
check("health carries ai block", "ai" in d and "ready" in d["ai"], d)
if FAKE:
    check("ai ready with model present", d["ai"]["ready"] is True, d["ai"])

# ============================================================ auth ============
section("Authentication")
s, d = login("admin", "root@localmind.test", "definitely-wrong")
check("wrong password -> INVALID_CREDENTIALS", s == 401 and err(d) == "INVALID_CREDENTIALS", d)
s, d = login("student", "root@localmind.test", INITIAL)
check("admin on student portal -> INVALID_CREDENTIALS", s == 401 and err(d) == "INVALID_CREDENTIALS", d)
s, d = login("wizard", "root@localmind.test", INITIAL)
check("unknown role -> INVALID_ROLE", s in (400, 404) and err(d) in ("INVALID_ROLE", "NOT_FOUND"), d)
s, d = call("GET", "/auth/me/")
check("unauthenticated -> 401", s == 401 and err(d) == "AUTHENTICATION_REQUIRED", d)
s, d = call("GET", "/auth/me/", tok="not.a.token")
check("garbage token -> 401", s == 401, d)

s, d = login("admin", "root@localmind.test", INITIAL)
if s == 200 and d.get("must_change_password"):
    tok = d["access"]
    s2, d2 = call("GET", "/admin/subjects/", tok=tok)
    check("must-change blocks normal endpoints", s2 == 403 and err(d2) == "PASSWORD_CHANGE_REQUIRED", d2)
    s2, d2 = call("POST", "/auth/password/change/", {"current_password": INITIAL, "new_password": "short"}, tok=tok)
    check("weak new password rejected", s2 == 400, d2)
    s2, d2 = call("POST", "/auth/password/change/", {"current_password": INITIAL, "new_password": INITIAL}, tok=tok)
    check("password reuse rejected", s2 == 400 and err(d2) in ("PASSWORD_REUSED", "VALIDATION_ERROR"), d2)
atok, arefresh, asess = login_and_settle("admin", "root@localmind.test")
s, d = call("GET", "/auth/me/", tok=atok)
check("me returns admin", s == 200 and d.get("role") == "admin", d)
s, d = call("POST", "/auth/password/change/", {"current_password": "nope", "new_password": STRONG + "x"}, tok=atok)
check("wrong current password rejected", s == 400 and err(d) == "INVALID_CURRENT_PASSWORD", d)

s, d = call("POST", "/auth/refresh/", {"refresh": arefresh})
check("refresh rotates tokens", s == 200 and d.get("access") and d.get("refresh") != arefresh, d)
new_refresh = d.get("refresh")
s, d = call("POST", "/auth/refresh/", {"refresh": arefresh})
check("old refresh token is blacklisted", s == 401 and err(d) == "INVALID_REFRESH", d)
arefresh = new_refresh
s, d = call("POST", "/auth/heartbeat/", {"session_id": asess}, tok=atok)
check("heartbeat", s == 200 and d.get("session_id"), d)

# ============================================================ admin ===========
section("Admin: subjects")
code = f"SYS{RUN}".upper()
s, d = call("POST", "/admin/subjects/", {"name": f"System Test {RUN}", "code": code.lower()}, tok=atok)
check("create subject (code uppercased)", s == 201 and d.get("code") == code, d)
subject = d["id"]
s, d = call("POST", "/admin/subjects/", {"name": "dup", "code": code}, tok=atok)
check("duplicate code -> SUBJECT_CODE_EXISTS", s in (400, 409) and err(d) == "SUBJECT_CODE_EXISTS", d)
s, d = call("POST", "/admin/subjects/", {"name": ""}, tok=atok)
check("missing fields -> VALIDATION_ERROR", s == 400 and err(d) == "VALIDATION_ERROR", d)
s, d = call("GET", f"/admin/subjects/{subject}/", tok=atok)
check("subject detail", s == 200 and d["id"] == subject, d)
s, d = call("GET", f"/admin/subjects/{uuid.uuid4()}/", tok=atok)
check("unknown subject -> 404", s == 404, d)
s, d = call("POST", f"/admin/subjects/{subject}/status/", {"status": "active"}, tok=atok)
check("unchanged status -> STATUS_UNCHANGED", s in (400, 409) and err(d) == "STATUS_UNCHANGED", d)

section("Admin: users")
fac_email = f"sys-fac-{RUN}@localmind.test"
s, d = call("POST", "/admin/faculty/", {"email": fac_email, "full_name": "Sys Faculty", "profile": {"employee_id": f"E{RUN}", "department": "CS"},
                                       "subject_ids": [subject]}, tok=atok)
check("create faculty with subject", s == 201 and d.get("role") == "faculty", d)
faculty_id = d["id"]
s, d = call("POST", "/admin/faculty/", {"email": fac_email, "full_name": "Dup"}, tok=atok)
check("duplicate email -> USER_EXISTS", s in (400, 409) and err(d) == "USER_EXISTS", d)
s, d = call("POST", "/admin/faculty/", {"email": "not-an-email", "full_name": "Bad"}, tok=atok)
check("invalid email rejected", s == 400, d)
s, d = call("GET", f"/admin/subjects/{subject}/", tok=atok)
check("faculty listed on subject detail", s == 200 and faculty_id in json.dumps(d), d)

stu_email = f"sys-stu-{RUN}@localmind.test"
s, d = call("POST", "/admin/students/", {"email": stu_email, "full_name": "Sys Student", "profile": {"roll_number": f"R{RUN}"}}, tok=atok)
check("create student", s == 201 and d.get("role") == "student", d)
student_id = d["id"]
s, d = call("POST", "/admin/students/", {"email": f"sys-stu2-{RUN}@localmind.test", "full_name": "Second Student"}, tok=atok)
student2_id = d.get("id")
s, d = call("GET", f"/admin/students/{student_id}/", tok=atok)
check("student detail", s == 200 and d["email"] == stu_email, d)
s, d = call("GET", "/admin/students/?q=Sys Student", tok=atok)
check("student list search", s == 200 and any(u["id"] == student_id for u in rows(d)), d)

s, d = call("POST", "/admin/students/import/", {}, tok=atok,
            files={"file": ("import.xlsx", make_xlsx(["Full Name", "E-mail", "Roll No"], [
                ["Imported One", f"sys-imp1-{RUN}@localmind.test", "I1"],
                ["Imported Two", stu_email, "I2"],
                ["Bad Row", "not-an-email", "I3"]]))})
check("excel import report", s in (200, 201) and d.get("created") == 1 and d.get("already_existing") == 1 and d.get("invalid") == 1, d)
s, d = call("POST", "/admin/students/import/", {}, tok=atok, files={"file": ("import.xlsx", make_xlsx(["Name"], [["x"]]))})
check("import missing headers -> MISSING_HEADERS", s == 400 and err(d) == "MISSING_HEADERS", d)
s, d = call("POST", "/admin/students/import/", {}, tok=atok, files={"file": ("import.txt", b"hello")})
check("import wrong file type rejected", s == 400, d)

section("Admin: enrollment and account state")
s, d = call("POST", f"/admin/subjects/{subject}/students/", {"student_ids": [student_id, str(uuid.uuid4())]}, tok=atok)
check("enroll with unknown id -> INVALID_STUDENT", s == 400 and err(d) == "INVALID_STUDENT", d)
s, d = call("POST", f"/admin/subjects/{subject}/students/", {"student_ids": [student_id, student2_id]}, tok=atok)
statuses = {r["student_id"]: r["status"] for r in d.get("results", [])} if s in (200, 201) else {}
check("enroll reports per-student status", statuses.get(student_id) == "enrolled" and statuses.get(student2_id) == "enrolled", d)
s, d = call("POST", f"/admin/subjects/{subject}/students/", {"student_ids": [student_id]}, tok=atok)
check("re-enroll active -> already_enrolled", any(r["status"] == "already_enrolled" for r in d.get("results", [])), d)
s, d = call("POST", f"/admin/subjects/{subject}/students/{student2_id}/discontinue/", tok=atok)
check("discontinue enrollment", s == 200, d)
s, d = call("POST", f"/admin/subjects/{subject}/students/{student2_id}/discontinue/", tok=atok)
check("discontinue twice -> ALREADY_DISCONTINUED", s in (400, 409) and err(d) == "ALREADY_DISCONTINUED", d)
s, d = call("POST", f"/admin/subjects/{subject}/students/", {"student_ids": [student2_id]}, tok=atok)
check("re-enroll discontinued -> re_enrolled", any(r["status"] == "re_enrolled" for r in d.get("results", [])), d)

s, d = call("POST", f"/admin/students/{student2_id}/discontinue/", {"reason": "left the programme"}, tok=atok)
check("discontinue user", s == 200 and d.get("status") == "discontinued", d)
s, d = login("student", f"sys-stu2-{RUN}@localmind.test", INITIAL)
check("discontinued user cannot log in", s == 401, d)
s, d = call("POST", f"/admin/students/{student2_id}/reactivate/", tok=atok)
check("reactivate user", s == 200 and d.get("status") == "active", d)
s, d = call("POST", f"/admin/students/{student2_id}/reset-password/", tok=atok)
check("reset password", s == 200, d)
s, d = login("student", f"sys-stu2-{RUN}@localmind.test", INITIAL)
check("reset restores initial password with forced change", s == 200 and d.get("must_change_password") is True, d)
s, d = call("GET", "/auth/me/", tok=atok)
s, d = call("POST", f"/admin/faculty/{d['id']}/discontinue/", tok=atok)
check("self-discontinue refused", s in (400, 403, 404, 409), d)

s, d = call("GET", f"/admin/audit-logs/?target_id={subject}", tok=atok)
check("audit log filtered by target", s == 200 and rows(d) and all(e.get("target_id") == subject for e in rows(d)), d)
s, d = call("GET", "/admin/audit-logs/?action=user.created", tok=atok)
check("audit log filtered by action", s == 200 and rows(d) and all(e["action"] == "user.created" for e in rows(d)), d)
s, d = call("GET", "/admin/analytics/platform/", tok=atok)
check("platform analytics", s == 200 and "users" in d and "quizzes" in d, d)
s, d = call("GET", "/admin/ai/status/?refresh=1", tok=atok)
check("admin ai status", s == 200 and "ready" in d, d)

# ============================================================ faculty ========
section("Faculty: scope and students")
ftok, frefresh, fsess = login_and_settle("faculty", fac_email)
s, d = call("GET", "/admin/subjects/", tok=ftok)
check("faculty on admin portal -> 403", s == 403, d)
s, d = call("GET", "/admin/ai/status/", tok=ftok)
check("faculty cannot read ai status", s == 403, d)
s, d = call("GET", "/faculty/subjects/", tok=ftok)
check("faculty sees assigned subject only", s == 200 and [x["id"] for x in rows(d)] == [subject], d)
s, d = call("GET", f"/faculty/subjects/{subject}/students/", tok=ftok)
check("faculty sees enrolled students", s == 200 and any(r["student_id"] == student_id for r in rows(d)), d)
s, d = call("GET", f"/faculty/students/search/?q=sys-imp1-{RUN}", tok=ftok)
check("faculty student search is minimal identity", s == 200 and len(rows(d)) == 1 and set(rows(d)[0]) <= {"id", "email", "full_name", "roll_number"}, d)
imported_id = rows(d)[0]["id"]
s, d = call("POST", f"/faculty/subjects/{subject}/students/", {"student_ids": [imported_id]}, tok=ftok)
check("faculty enrolls student", s in (200, 201), d)

# a second faculty who must see nothing of this subject
other_email = f"sys-other-{RUN}@localmind.test"
call("POST", "/admin/faculty/", {"email": other_email, "full_name": "Other Faculty"}, tok=atok)
otok, _, _ = login_and_settle("faculty", other_email)
s, d = call("GET", f"/faculty/subjects/{subject}/students/", tok=otok)
check("unassigned faculty -> 404 on subject", s == 404, d)

section("Faculty: documents")
s, d = call("POST", "/faculty/documents/", {"subject_id": subject}, tok=ftok, files={"file": ("notes.txt", b"plain text")})
check("unsupported file type rejected", s == 400 and err(d) == "UNSUPPORTED_FILE_TYPE", d)
s, d = call("POST", "/faculty/documents/", {"subject_id": subject}, tok=ftok, files={"file": ("fake.pdf", b"not really a pdf")})
check("content/extension mismatch rejected", s == 400 and err(d) == "FILE_CONTENT_MISMATCH", d)
s, d = call("POST", "/faculty/documents/", {"subject_id": subject}, tok=ftok, files={"file": ("empty.docx", b"")})
check("empty file rejected", s == 400, d)
s, d = call("POST", "/faculty/documents/", {"subject_id": subject, "title": f"SYS Book {RUN}"}, tok=ftok, files={"file": ("../evil.docx", make_docx())})
check("docx upload", s == 201 and d.get("status") == "uploaded", d)
doc = d["id"]
check("stored name is not the client path", "evil" not in json.dumps(d.get("file", "")) and ".." not in json.dumps(d), d)
s, d = call("POST", f"/faculty/documents/{doc}/", tok=otok)
s, d = call("GET", f"/faculty/documents/{doc}/", tok=otok)
check("other faculty cannot see document", s == 404, d)
s, d = call("POST", f"/faculty/documents/{doc}/process/", tok=ftok)
check("process accepted", s in (200, 202) and d.get("status") in ("processing", "under_review"), d)
for _ in range(60):
    s, d = call("GET", f"/faculty/documents/{doc}/", tok=ftok)
    if d.get("status") in ("under_review", "error"):
        break
    time.sleep(1)
check("processing finished", d.get("status") == "under_review", d.get("error_message") or d.get("status"))
check("outline came from the AI", d.get("outline_source") == "ai" if FAKE else d.get("outline_source") in ("ai", "source_hierarchy"), d.get("outline_source"))

s, d = call("GET", f"/faculty/documents/{doc}/outline/", tok=ftok)
check("outline has chapters, modules and headings", s == 200 and len(d["chapters"]) == 2 and d.get("headings"), d.get("chapters") and len(d["chapters"]))
outline = d
mods = [m for c in outline["chapters"] for m in c["modules"]]
check("modules carry source text", all(m.get("source_text") for m in mods), [m["title"] for m in mods])
first_mod = mods[0]
# edit: rename first module, delete last module, keep ids
edited = {"document_title": outline.get("document_title"), "chapters": []}
for c in outline["chapters"]:
    edited["chapters"].append({"id": c["id"], "title": c["title"], "source_heading_index": c.get("source_heading_index"),
                               "modules": [{"id": m["id"], "title": m["title"], "source_heading_index": m.get("source_heading_index")} for m in c["modules"]]})
edited["chapters"][0]["modules"][0]["title"] = "Process Management (edited)"
removed = edited["chapters"][0]["modules"].pop()
s, d = call("PUT", f"/faculty/documents/{doc}/outline/", edited, tok=ftok)
check("outline PUT edits in place and deletes omitted module", s == 200, d)
s, d = call("GET", f"/faculty/documents/{doc}/outline/", tok=ftok)
titles = [m["title"] for c in d["chapters"] for m in c["modules"]]
check("rename persisted, deleted module gone", "Process Management (edited)" in titles and removed["title"] not in titles, titles)
s, d = call("PUT", f"/faculty/documents/{doc}/outline/", {"chapters": []}, tok=ftok)
check("empty outline rejected", s == 400 and err(d) == "EMPTY_OUTLINE", d)
s, d = call("PUT", f"/faculty/documents/{doc}/outline/", {"chapters": [{"title": "", "modules": []}]}, tok=ftok)
check("chapter without title rejected", s == 400 and err(d) in ("MISSING_TITLE", "VALIDATION_ERROR"), d)

s, d = call("GET", f"/faculty/documents/{doc}/", tok=ftok)
version_before = d.get("content_version")
s, d = call("PATCH", f"/faculty/modules/{first_mod['id']}/", {"source_text": first_mod["source_text"] + "\n\nAn added sentence for the tutor."}, tok=ftok)
check("module source edit", s == 200, d)
s, d = call("GET", f"/faculty/documents/{doc}/", tok=ftok)
check("content_version bumped by edit", d.get("content_version", 0) > (version_before or 0), (version_before, d.get("content_version")))

s, d = call("POST", f"/faculty/documents/{doc}/ready/", tok=ftok)
check("mark ready", s == 200 and d.get("status") == "ready", d)
s, d = call("POST", f"/faculty/documents/{doc}/ready/", tok=ftok)
check("ready twice -> INVALID_STATE", s in (400, 409) and err(d) == "INVALID_STATE", d)
s, d = call("POST", f"/faculty/documents/{doc}/publish/", tok=ftok)
check("publish", s == 200 and d.get("status") == "published", d)
s, d = call("GET", f"/faculty/documents/{doc}/outline/", tok=ftok)
mods = [m for c in d["chapters"] for m in c["modules"]]
chapter1, chapter2 = d["chapters"][0], d["chapters"][1]
check("modules with source opened on publish", all(m["availability"] == "open" for m in mods), [m["availability"] for m in mods])
s, d = call("POST", f"/faculty/chapters/{chapter2['id']}/availability/", {"availability": "locked"}, tok=ftok)
check("lock second chapter after publish", s == 200 and all(m["availability"] == "locked" for m in d), d)
s, d = call("POST", f"/faculty/modules/{mods[0]['id']}/availability/", {"availability": "locked"}, tok=ftok)
check("lock first module", s == 200 and d.get("availability") == "locked", d)
s, d = call("POST", f"/faculty/modules/{mods[0]['id']}/availability/", {"availability": "open"}, tok=ftok)
check("open first module", s == 200 and d.get("availability") == "open", d)
s, d = call("POST", f"/faculty/modules/{mods[0]['id']}/availability/", {"availability": "sideways"}, tok=ftok)
check("invalid availability rejected", s == 400, d)
s, d = call("POST", f"/faculty/chapters/{chapter1['id']}/availability/", {"availability": "open"}, tok=ftok)
check("open whole chapter", s == 200, d)
open_mod = mods[0]
locked_mod = chapter2["modules"][0]

section("Faculty: quizzes")
s, d = call("POST", "/faculty/quizzes/", {"title": "bad", "questions": []}, tok=ftok)
check("quiz without target -> TARGET_REQUIRED", s == 400 and err(d) == "TARGET_REQUIRED", d)
s, d = call("POST", "/faculty/quizzes/", {"module_id": open_mod["id"], "title": "bad",
                                          "questions": [{"type": "mcq", "question": "Q", "options": [{"key": "A", "text": "x"}], "correct_answer": "A"}]}, tok=ftok)
check("invalid mcq -> INVALID_QUESTIONS", s == 400 and err(d) == "INVALID_QUESTIONS", d)
manual_q = [{"type": "mcq", "question": "Which is a scheduling policy?", "options": [{"key": "A", "text": "Round robin"}, {"key": "B", "text": "Paging"}, {"key": "C", "text": "TCP"}, {"key": "D", "text": "Segmentation"}],
             "correct_answer": "A", "explanation": "Round robin gives each process a time slice.", "source_reference": "Round robin gives each process a time slice"},
            {"type": "subjective", "question": "Explain round robin scheduling.", "expected_rubric": "Each process gets a time slice; processes rotate.", "source_reference": "Round robin"}]
s, d = call("POST", "/faculty/quizzes/", {"module_id": open_mod["id"], "title": "SYS manual quiz", "questions": manual_q, "pass_percentage": 50, "max_attempts": 2}, tok=ftok)
check("manual quiz created", s == 201 and d.get("generator") == "manual" and d["status"] == "draft", d)
manual_quiz = d["id"]

s, d = call("POST", "/faculty/quizzes/generate/", {"module_id": open_mod["id"], "num_mcqs": 3, "num_subjective": 1}, tok=ftok)
check("AI quiz generated", s == 201 and d.get("generator") == ("ai" if FAKE else d.get("generator")) and len(d["questions"]) == 4, d)
ai_quiz = d["id"]
check("generated questions have distinct A-D options", all(len({o["text"] for o in q["options"]}) == 4 for q in d["questions"] if q["type"] == "mcq"), d["questions"][0])
s, d = call("POST", "/faculty/quizzes/generate/", {"chapter_id": chapter1["id"], "num_mcqs": 2}, tok=ftok)
check("chapter quiz generated", s == 201 and d.get("kind") == "chapter", d)
s, d = call("POST", "/faculty/quizzes/generate/", {"module_id": open_mod["id"], "num_mcqs": 0, "num_subjective": 0}, tok=ftok)
check("zero counts -> INVALID_COUNTS", s == 400 and err(d) == "INVALID_COUNTS", d)

if FAKE:
    fake_control(fail_next=2)  # one call + one retry both bad -> fallback
    s, d = call("POST", "/faculty/quizzes/generate/", {"module_id": open_mod["id"], "num_mcqs": 2}, tok=ftok)
    check("bad model output twice -> fallback draft", s == 201 and d.get("generator") == "fallback" and d.get("generation_warning"), d)
    fb_quiz = d["id"]
    s, d = call("POST", f"/faculty/quizzes/{fb_quiz}/status/", {"status": "published"}, tok=ftok)
    check("placeholder quiz cannot be published", s in (400, 409) and err(d) == "PLACEHOLDER_QUESTIONS", d)
    fake_control(fail_next=1)
    s, d = call("POST", "/faculty/quizzes/generate/", {"module_id": open_mod["id"], "num_mcqs": 2}, tok=ftok)
    check("bad output once is recovered by the retry", s == 201 and d.get("generator") == "ai", d)

s, d = call("POST", f"/faculty/quizzes/{manual_quiz}/status/", {"status": "published"}, tok=ftok)
check("publish manual quiz", s == 200 and d["status"] == "published", d)
s, d = call("POST", f"/faculty/quizzes/{ai_quiz}/status/", {"status": "published"}, tok=ftok)
check("publish ai quiz", s == 200 and d["status"] == "published", d)
s, d = call("GET", f"/faculty/quizzes/?subject={subject}", tok=ftok)
check("quiz list", s == 200 and len(rows(d)) >= 3, d)
s, d = call("GET", f"/faculty/quizzes/{manual_quiz}/", tok=otok)
check("other faculty cannot see quiz", s == 404, d)

section("Faculty: assignments")
s, d = call("POST", "/faculty/assignments/", {"module_id": open_mod["id"], "title": "bad", "max_score": 10, "rubric": [{"criterion": "a", "points": 3}]}, tok=ftok)
check("rubric not summing to max_score rejected", s == 400, d)
s, d = call("POST", "/faculty/assignments/", {"module_id": open_mod["id"], "title": "SYS essay", "max_score": 10,
                                              "rubric": [{"criterion": "Accuracy", "points": 6}, {"criterion": "Clarity", "points": 4}], "allow_resubmission": False}, tok=ftok)
check("create assignment", s == 201 and d["status"] == "draft", d)
assignment = d["id"]
s, d = call("POST", "/faculty/assignments/generate/", {"module_id": open_mod["id"], "max_score": 20, "focus": "scheduling"}, tok=ftok)
check("AI assignment generated with rubric summing to max", s == 201 and sum(r["points"] for r in d["rubric"]) == 20 and (d["generator"] == "ai" if FAKE else True), d)
s, d = call("POST", f"/faculty/assignments/{assignment}/status/", {"status": "published"}, tok=ftok)
check("publish assignment", s == 200 and d["status"] == "published", d)

# ============================================================ student ========
section("Student: access and reading")
stok, srefresh, ssess = login_and_settle("student", stu_email)
s, d = call("GET", "/faculty/subjects/", tok=stok)
check("student on faculty portal -> 403", s == 403, d)
s, d = call("GET", "/student/subjects/", tok=stok)
check("student sees enrolled subject", s == 200 and any(x["id"] == subject for x in rows(d)), d)
s, d = call("GET", f"/student/subjects/{subject}/documents/", tok=stok)
check("published book listed with progress fields", s == 200 and rows(d) and "progress_percent" in rows(d)[0], d)
s, d = call("GET", f"/student/documents/{doc}/", tok=stok)
smods = [m for c in d["chapters"] for m in c["modules"]]
check("document tree without source text", s == 200 and smods and all("source_text" not in m for m in smods), d)
s, d = call("GET", f"/student/modules/{locked_mod['id']}/", tok=stok)
check("locked module -> MODULE_LOCKED", s == 403 and err(d) == "MODULE_LOCKED", d)
s, d = call("GET", f"/student/modules/{open_mod['id']}/", tok=stok)
check("open module returns source and marks in_progress", s == 200 and d.get("source_text") and (d.get("progress") or {}).get("status") == "in_progress", d)
s, d = call("POST", f"/student/modules/{open_mod['id']}/time/", {"seconds": 5000}, tok=stok)
check("time chunk clamped to 15 minutes", s == 200 and d.get("learning_seconds") == 900, d)
s, d = call("POST", f"/student/modules/{open_mod['id']}/time/", {"seconds": -5}, tok=stok)
check("negative time rejected", s == 400 and err(d) == "VALIDATION_ERROR", d)
s, d = call("POST", f"/student/modules/{open_mod['id']}/time/", {"seconds": "lots"}, tok=stok)
check("non-numeric time rejected", s == 400, d)
s, d = call("POST", "/auth/heartbeat/", {"session_id": ssess}, tok=stok)
check("student heartbeat", s == 200, d)

section("Student: tutor")
s, d = call("POST", f"/student/modules/{open_mod['id']}/teach/", tok=stok)
check("teach returns structured lesson", s == 200 and d["lesson"].get("sections") and d.get("generator") == ("ai" if FAKE else d.get("generator")), d)
check("teach not cached on first call", d.get("cached") is False, d)
s, d = call("POST", f"/student/modules/{open_mod['id']}/teach/", tok=stok)
check("teach cached on second call", s == 200 and d.get("cached") is True, d)
s, d = call("POST", f"/student/modules/{locked_mod['id']}/teach/", tok=stok)
check("teach on locked module refused", s == 403, d)
s, d = call("POST", f"/student/modules/{open_mod['id']}/ask/", {"question": ""}, tok=stok)
check("empty question rejected", s == 400, d)
s, d = call("POST", f"/student/modules/{open_mod['id']}/ask/", {"question": "x" * 2001}, tok=stok)
check("over-long question rejected", s == 400, d)
s, d = call("POST", f"/student/modules/{open_mod['id']}/ask/", {"question": "What is a process?"}, tok=stok)
ask_ok = check("ask answers and opens a conversation", s in (200, 201) and (d.get("message") or {}).get("content") and d.get("conversation_id") and d.get("follow_up_suggestions") is not None, d)
conv = d.get("conversation_id")
if ask_ok:
    s, d = call("POST", f"/student/modules/{open_mod['id']}/ask/", {"question": "And a thread?", "conversation_id": conv}, tok=stok)
    check("ask continues conversation", s in (200, 201) and d.get("conversation_id") == conv, d)
    s, d = call("GET", f"/student/conversations/{conv}/", tok=stok)
    check("conversation has 4 messages", s == 200 and len(d.get("messages", [])) == 4, d)
    s, d = call("POST", f"/student/modules/{locked_mod['id']}/ask/", {"question": "Q", "conversation_id": conv}, tok=stok)
    check("conversation on other module refused", s in (400, 403), d)
s, d = call("GET", f"/student/conversations/?module={open_mod['id']}", tok=stok)
check("conversation list", s == 200 and len(rows(d)) >= 1, d)

section("Student: quizzes")
s, d = call("GET", f"/student/quizzes/?module={open_mod['id']}", tok=stok)
check("published quizzes on open module listed", s == 200 and {q["id"] for q in rows(d)} >= {manual_quiz, ai_quiz}, d)
s, d = call("POST", f"/student/quizzes/{manual_quiz}/attempts/", tok=stok)
check("start attempt hides answers and rubrics", s == 201 and d.get("attempt_id") and all("correct_answer" not in q and "expected_rubric" not in q for q in d["questions"]), d)
attempt = d["attempt_id"]
s, d2 = call("POST", f"/student/quizzes/{manual_quiz}/attempts/", tok=stok)
check("start again resumes same attempt", s in (200, 201) and d2.get("attempt_id") == attempt, d2)
s, d = call("POST", f"/student/quiz-attempts/{attempt}/submit/", {"submitted_answers": {"q1": "A", "q2": "Each process gets a time slice and they rotate; this is correct."}}, tok=stok)
check("submit scores mcq deterministically and evaluates subjective", s == 200 and d.get("status") == "evaluated" and d["percentage"] == 100.0 and d["passed"] is True, d)
check("time is server computed", isinstance(d.get("time_taken_seconds"), int), d)
s, d = call("POST", f"/student/quiz-attempts/{attempt}/submit/", {"submitted_answers": {"q1": "A"}}, tok=stok)
check("resubmit -> ALREADY_SUBMITTED", s == 409 and err(d) == "ALREADY_SUBMITTED", d)
s, d = call("GET", f"/student/quiz-attempts/{attempt}/", tok=stok)
check("attempt detail visible to owner", s == 200 and d["id"] == attempt, d)
s, d = call("POST", f"/student/quizzes/{manual_quiz}/attempts/", tok=stok)
attempt2 = d.get("attempt_id")
s, d = call("POST", f"/student/quiz-attempts/{attempt2}/submit/", {"submitted_answers": {"q1": "B", "q2": ""}}, tok=stok)
check("second attempt wrong answers scored 0", s == 200 and d["percentage"] == 0.0 and d["passed"] is False, d)
s, d = call("POST", f"/student/quizzes/{manual_quiz}/attempts/", tok=stok)
check("third attempt -> MAX_ATTEMPTS_REACHED", s in (400, 409) and err(d) == "MAX_ATTEMPTS_REACHED", d)
s, d = call("POST", f"/student/quiz-attempts/{attempt2}/remediation/", tok=stok)
check("remediation for missed questions", s == 200 and len(d.get("items", [])) == 2, d)
s, d = call("POST", f"/student/quiz-attempts/{attempt}/remediation/", tok=stok)
check("remediation with nothing wrong", s == 200 and d.get("items") == [], d)
s, d = call("GET", f"/student/scores/?subject={subject}", tok=stok)
check("scores list", s == 200 and len(rows(d)) == 2, d)
s, d = call("GET", "/student/analytics/overview/", tok=stok)
check("student overview", s == 200 and d.get("quizzes", {}).get("attempts", 0) >= 2 and d.get("time", {}).get("learning_seconds", 0) >= 900, d)

# a second student must not see the first student's attempt
s2tok, _, _ = login_and_settle("student", f"sys-imp1-{RUN}@localmind.test")
s, d = call("GET", f"/student/quiz-attempts/{attempt}/", tok=s2tok)
check("other student cannot read attempt", s == 404, d)
s, d = call("POST", f"/student/quiz-attempts/{attempt}/submit/", {"submitted_answers": {}}, tok=s2tok)
check("other student cannot submit attempt", s in (403, 404), d)

section("Student: assignments")
s, d = call("GET", "/student/assignments/", tok=stok)
check("published assignment listed", s == 200 and any(a["id"] == assignment for a in rows(d)), d)
s, d = call("POST", f"/student/assignments/{assignment}/submissions/", {"content": ""}, tok=stok)
check("empty submission rejected", s == 400, d)
s, d = call("POST", f"/student/assignments/{assignment}/submissions/", {"content": "Round robin gives each process a time slice.", "time_spent_seconds": 300}, tok=stok)
check("assignment submitted", s == 201 and d.get("status") in ("submitted", "pending", "pending_evaluation"), d)
submission = d["id"]
s, d = call("POST", f"/student/assignments/{assignment}/submissions/", {"content": "again"}, tok=stok)
check("resubmission refused when not allowed", s in (400, 409), d)
s, d = call("GET", "/student/assignment-submissions/", tok=stok)
check("own submissions listed", s == 200 and any(x["id"] == submission for x in rows(d)), d)

# ============================================================ faculty review =
section("Faculty: review and analytics")
s, d = call("GET", f"/faculty/quizzes/{manual_quiz}/attempts/", tok=ftok)
check("faculty sees attempts with detailed results", s == 200 and len(rows(d)) == 2 and rows(d)[0].get("detailed_results"), d)
s, d = call("POST", f"/faculty/quiz-attempts/{attempt2}/re-evaluate/", {"overrides": {"q2": {"score_awarded": 1.0, "feedback": "Partial credit."}}}, tok=ftok)
check("faculty override re-scores", s == 200 and d.get("percentage") == 50.0, d)
s, d = call("PATCH", f"/faculty/quizzes/{manual_quiz}/", {"title": "SYS manual quiz v2"}, tok=ftok)
check("metadata edit keeps the same version", s == 200 and d["id"] == manual_quiz and d.get("title") == "SYS manual quiz v2", d)
s, d = call("PATCH", f"/faculty/quizzes/{manual_quiz}/", {"questions": manual_q[:1]}, tok=ftok)
check("question edit with attempts creates new version", s == 200 and d["id"] != manual_quiz and d.get("version") == 2 and d["status"] == "draft", d)
s, d = call("GET", f"/faculty/quizzes/{manual_quiz}/", tok=ftok)
check("old version superseded", s == 200 and d["status"] == "superseded", d)
s, d = call("GET", f"/faculty/assignments/{assignment}/submissions/", tok=ftok)
check("faculty sees submission", s == 200 and any(x["id"] == submission for x in rows(d)), d)
s, d = call("POST", f"/faculty/assignment-submissions/{submission}/evaluate/", {"score": 99, "feedback": "x"}, tok=ftok)
check("score above max rejected", s == 400, d)
s, d = call("POST", f"/faculty/assignment-submissions/{submission}/evaluate/", {"score": 8, "feedback": "Good work."}, tok=ftok)
check("evaluate submission", s == 200 and d.get("status") == "evaluated" and d.get("score") == 8, d)
s, d = call("GET", "/student/assignment-submissions/", tok=stok)
check("student sees score and feedback", any(x["id"] == submission and x.get("score") == 8 for x in rows(d)), d)
s, d = call("GET", "/faculty/analytics/overview/", tok=ftok)
check("faculty overview", s == 200, d)
s, d = call("GET", f"/faculty/analytics/subjects/{subject}/", tok=ftok)
check("subject analytics", s == 200 and d.get("students_enrolled", 0) >= 3, d)
s, d = call("GET", f"/faculty/analytics/subjects/{subject}/students/", tok=ftok)
row = next((r for r in d.get("students", []) if r["student_id"] == student_id), None)
check("cohort row reflects quiz, time and assignment", row and row.get("learning_seconds", 0) >= 900 and row.get("assignment_average") is not None, row)
s, d = call("GET", f"/faculty/analytics/subjects/{subject}/modules/", tok=ftok)
check("module funnel", s == 200 and d.get("modules"), d)
s, d = call("GET", f"/faculty/analytics/students/{student_id}/", tok=ftok)
check("student analytics for shared student", s == 200, d)
s, d = call("GET", f"/faculty/analytics/students/{student_id}/", tok=otok)
check("student analytics refused without shared subject", s == 404, d)

section("Lifecycle: unpublish, archive, structure lock")
s, d = call("PUT", f"/faculty/documents/{doc}/outline/", edited, tok=ftok)
check("published structure locked", s in (400, 409) and err(d) == "PUBLISHED_STRUCTURE_LOCKED", d)
s, d = call("POST", f"/faculty/documents/{doc}/unpublish/", tok=ftok)
check("unpublish", s == 200 and d["status"] == "unpublished", d)
s, d = call("GET", f"/student/subjects/{subject}/documents/", tok=stok)
check("unpublished book hidden from student", s == 200 and not any(x["id"] == doc for x in rows(d)), d)
s, d = call("GET", f"/student/modules/{open_mod['id']}/", tok=stok)
check("module of unpublished book refused", s in (403, 404), d)
s, d = call("POST", f"/faculty/documents/{doc}/publish/", tok=ftok)
check("republish", s == 200 and d["status"] == "published", d)
s, d = call("POST", f"/faculty/quizzes/{ai_quiz}/status/", {"status": "closed"}, tok=ftok)
check("close quiz", s == 200 and d["status"] == "closed", d)
s, d = call("POST", f"/student/quizzes/{ai_quiz}/attempts/", tok=stok)
check("closed quiz refuses attempts", s in (400, 404, 409), d)

# ============================================================ logout =========
section("Logout and revocation")
s, d = call("POST", "/auth/logout/", {"refresh": srefresh, "session_id": ssess}, tok=stok)
check("student logout", s in (200, 204), d)
s, d = call("POST", "/auth/refresh/", {"refresh": srefresh})
check("refresh after logout refused", s == 401, d)
s, d = call("GET", f"/admin/analytics/users/{student_id}/sessions/", tok=atok)
sess_rows = d.get("sessions", []) if s == 200 else []
check("admin sees closed session with logout reason", any(r.get("ended_by") == "logout" for r in sess_rows), sess_rows[:2])
s, d = call("POST", f"/admin/faculty/{faculty_id}/discontinue/", tok=atok)
s, d = call("GET", "/faculty/subjects/", tok=ftok)
check("discontinued faculty access token rejected", s in (401, 403), d)
s, d = call("POST", "/auth/refresh/", {"refresh": frefresh})
check("discontinued faculty refresh revoked", s == 401, d)
call("POST", f"/admin/faculty/{faculty_id}/reactivate/", tok=atok)

# ============================================================ AI outage ======
section("Concurrency and pagination")
import threading
s, d = call("POST", "/faculty/documents/", {"subject_id": subject, "title": "SYS concurrent"}, tok=ftok, files={"file": ("book2.docx", make_docx())})
doc2 = d["id"]
codes = []
def _proc():
    codes.append(call("POST", f"/faculty/documents/{doc2}/process/", tok=ftok)[0])
threads = [threading.Thread(target=_proc) for _ in range(4)]
[t.start() for t in threads]; [t.join() for t in threads]
check("parallel process calls: exactly one wins", sorted(codes) == [200, 409, 409, 409] or sorted(codes) == [202, 409, 409, 409], codes)
s, d = call("GET", "/admin/audit-logs/?page_size=2", tok=atok)
check("page_size honoured", s == 200 and len(d.get("results", [])) == 2 and d.get("next"), {k: d.get(k) for k in ("count", "next")})
s, d = call("GET", "/admin/audit-logs/?page_size=5000", tok=atok)
check("page_size capped at 200", s == 200 and len(d.get("results", [])) <= 200, len(d.get("results", [])))
s, d = call("GET", "/admin/audit-logs/?page=99999", tok=atok)
check("out-of-range page -> 404", s == 404, d)
s, d = call("GET", "/faculty/subjects/", tok=atok)
check("admin may use faculty portal (unscoped)", s == 200 and any(x["id"] == subject for x in rows(d)), d)
s, d = call("POST", f"/student/modules/{locked_mod['id']}/time/", {"seconds": 10}, tok=s2tok)
check("time on locked module refused", s == 403, d)

if FAKE:
    section("AI outage behaviour (live)")
    fake_control(offline=1)
    s, d = call("GET", "/admin/ai/status/?refresh=1", tok=atok)
    check("ai status reports unreachable", s == 200 and d.get("ready") is False and d.get("reachable") is False, d)
    s, d = call("GET", "/health/")
    check("health stays ok while ai is down", s == 200 and d["status"] == "ok" and d["ai"]["ready"] is False, d)
    s, d = call("POST", f"/student/modules/{open_mod['id']}/ask/", {"question": "Anything?"}, tok=s2tok)
    check("ask -> 503 AI_UNAVAILABLE with reason", s == 503 and err(d) == "AI_UNAVAILABLE" and d["error"]["details"].get("reason") == "unavailable", d)
    s, d = call("GET", f"/student/conversations/", tok=s2tok)
    check("failed question is still recorded in the thread", s == 200 and rows(d), d)
    s, d = call("POST", f"/student/modules/{open_mod['id']}/teach/", tok=s2tok)
    check("teach serves cached ai lesson while down", s == 200 and d.get("generator") == "ai" and d.get("cached") is True, d)
    s, d = call("POST", f"/student/modules/{chapter1['modules'][1]['id']}/teach/", tok=s2tok)
    check("uncached teach falls back to source lesson", s == 200 and d.get("generator") == "fallback" and d.get("ai_error") == "unavailable", d)
    s, d = call("POST", "/faculty/quizzes/generate/", {"module_id": open_mod["id"], "num_mcqs": 2}, tok=ftok)
    check("quiz generation falls back with warning", s == 201 and d.get("generator") == "fallback" and "unavailable" in (d.get("generation_warning") or ""), d)
    s, d = call("POST", "/faculty/assignments/generate/", {"module_id": open_mod["id"], "max_score": 10}, tok=ftok)
    check("assignment generation falls back", s == 201 and d.get("generator") == "fallback", d)
    s, d = call("POST", f"/student/quiz-attempts/{attempt2}/remediation/", tok=stok)
    check("remediation falls back to source passages", s == 200 and d.get("generator") == "fallback" and d.get("items"), d)
    # subjective grading cannot run: attempt goes pending, faculty override resolves it
    s, d = call("POST", "/faculty/quizzes/", {"module_id": open_mod["id"], "title": "SYS pending quiz", "questions": manual_q[1:]}, tok=ftok)
    pend_quiz = d["id"]
    call("POST", f"/faculty/quizzes/{pend_quiz}/status/", {"status": "published"}, tok=ftok)
    s, d = call("POST", f"/student/quizzes/{pend_quiz}/attempts/", tok=s2tok)
    check("second student starts pending quiz", s == 201 and d.get("attempt_id"), d)
    pend_attempt = d.get("attempt_id")
    s, d = call("POST", f"/student/quiz-attempts/{pend_attempt}/submit/", {"submitted_answers": {"q1": "time slices rotate"}}, tok=s2tok)
    check("subjective submit while ai down -> pending_evaluation", s == 200 and d.get("status") == "pending_evaluation", d)
    s, d = call("POST", f"/faculty/quiz-attempts/{pend_attempt}/re-evaluate/", {"overrides": {"q1": {"score_awarded": 1.0, "feedback": "ok"}}}, tok=ftok)
    check("faculty override resolves pending attempt", s == 200 and d.get("status") == "evaluated" and d.get("percentage") == 100.0, d)
    fake_control(offline=0)
    s, d = call("GET", "/admin/ai/status/?refresh=1", tok=atok)
    check("ai status recovers", s == 200 and d.get("ready") is True, d)
    s, d = call("POST", f"/faculty/quiz-attempts/{pend_attempt}/re-evaluate/", {}, tok=ftok)
    check("re-evaluate with ai back re-runs grading", s == 200 and d.get("status") == "evaluated", d)

# ============================================================ throttling =====
section("Rate limiting (login, run last)")
codes = [login("student", "nobody@localmind.test", "x")[0] for _ in range(25)]
check("login throttled after 20/min", 429 in codes, codes)

fake_control(offline=0, fail_next=0)
print(f"\n{passed} passed, {failed} failed")
sys.exit(failed)
